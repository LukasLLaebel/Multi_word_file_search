import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document

# Funktion til at søge efter et ord i et dokument og returnere linjenummer og tekst
def search_word_in_docx(file_path, search_word):
    doc = Document(file_path)
    found_lines = []
    for idx, para in enumerate(doc.paragraphs):
        if search_word.lower() in para.text.lower():
            found_lines.append((idx + 1, para.text))
    return found_lines

# Funktion til at søge i alle Word-dokumenter i en mappe og undermapper
def search_in_folder(folder_path, search_word):
    results = []
    for root, dirs, files in os.walk(folder_path):
        for filename in files:
            if filename.endswith(".docx") and not filename.startswith("~$"):  # Undgå midlertidige filer
                file_path = os.path.join(root, filename)
                try:
                    # Åben dokumentet og kontroller at det er gyldigt
                    doc = Document(file_path)
                    found_lines = search_word_in_docx(file_path, search_word)
                    if found_lines:
                        results.append((filename, file_path, found_lines))
                except Exception as e:
                    print(f"Fejl ved behandling af fil: {file_path}. Fejl: {str(e)}")
    return results

# Funktion til at åbne en fil
def open_file(file_path):
    try:
        # Brug os.startfile til at åbne filen med standardprogrammet (som bør være Word for .docx)
        os.startfile(file_path)  # Dette virker kun på Windows
    except Exception as e:
        messagebox.showerror("Fejl", f"Kunne ikke åbne filen: {str(e)}")

# Funktion til at starte søgningen
def start_search():
    global search_results
    folder_path = filedialog.askdirectory()
    search_word = entry_search.get()
    
    if not search_word:
        messagebox.showwarning("Advarsel", "Indtast et søgeord!")
        return
    
    if not folder_path:
        messagebox.showwarning("Advarsel", "Vælg en mappe!")
        return
    
    search_results = search_in_folder(folder_path, search_word)
    
    if search_results:
        text_results.config(state=tk.NORMAL)  # Tillad redigering af tekstfeltet midlertidigt
        text_results.delete(1.0, tk.END)  # Ryd tekstfeltet
        for filename, file_path, found_lines in search_results:
            # Tilføj overskrift for filnavnet
            text_results.insert(tk.END, f"\nFil: {filename}\n", "bold")
            
            # Tilføj hver linje hvor ordet findes
            for line_num, line_text in found_lines:
                text_results.insert(tk.END, f"  Linje {line_num}: ", "bold")
                text_results.insert(tk.END, f"{line_text[:50]}...\n")
        text_results.config(state=tk.DISABLED)  # Gør tekstfeltet skrivebeskyttet igen
    else:
        messagebox.showinfo("Ingen resultater", "Ingen resultater fundet.")

# Funktion til at håndtere klik på tekstfeltet
def on_text_click(event):
    # Find den linje der blev klikket på
    index = text_results.index(tk.CURRENT)
    clicked_line = text_results.get(f"{index} linestart", f"{index} lineend")
    
    # Find den fil, som er associeret med den klikkede linje
    for filename, file_path, found_lines in search_results:
        if filename in clicked_line:
            open_file(file_path)

# Opret GUI
root = tk.Tk()
root.title("Word Dokument Søgemaskine")

# Label og input til søgeord
label_search = tk.Label(root, text="Indtast søgeord:")
label_search.pack(pady=5)
entry_search = tk.Entry(root, width=50)
entry_search.pack(pady=5)

# Knappen til at starte søgningen
btn_search = tk.Button(root, text="Søg i Word dokumenter", command=start_search)
btn_search.pack(pady=5)

# Tekstboks til resultater
text_results = tk.Text(root, width=100, height=20, wrap="word")
text_results.pack(pady=10)

# Definer formateringsstile for tekstboks
text_results.tag_configure("bold", font=("Helvetica", 10, "bold"))

# Gør tekstfeltet skrivebeskyttet
text_results.config(state=tk.DISABLED)

# Håndtering af klik i tekstfeltet
text_results.bind("<Button-1>", on_text_click)

# Variabel til at holde søgeresultater
search_results = []

# Start GUI'en
root.mainloop()
